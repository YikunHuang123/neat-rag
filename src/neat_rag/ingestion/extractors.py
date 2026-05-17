import time
from pathlib import Path
from typing import Dict, Any, Tuple

from neat_rag.exceptions import ExtractionError, UnsupportedFileTypeError
from neat_rag.logger import get_logger

logger = get_logger(__name__)


class PdfExtractor:
    """Extracts text and structure from PDF files using Docling."""

    def __init__(
        self,
        enable_ocr: bool = False,
        include_images: bool = True,
        include_tables: bool = True,
        images_scale: float = 2.0,
    ):
        self._enable_ocr = enable_ocr
        self._include_images = include_images
        self._include_tables = include_tables
        self._images_scale = images_scale
        self._converter = None  # lazy init — docling is slow to import

    def _get_converter(self):
        if self._converter is None:
            from docling.document_converter import DocumentConverter, PdfFormatOption
            from docling.datamodel.base_models import InputFormat
            from docling.datamodel.pipeline_options import PdfPipelineOptions

            opts = PdfPipelineOptions()
            opts.do_ocr = self._enable_ocr
            opts.do_picture_description = self._include_images
            opts.do_table_structure = self._include_tables
            opts.images_scale = self._images_scale

            if self._include_images or self._enable_ocr:
                logger.info(
                    "🚀 Initializing Docling models... If this is the first run, it may download ~1GB of weights (SmolVLM/OCR). "
                    "This might take several minutes depending on your network speed.",
                    include_images=self._include_images,
                    enable_ocr=self._enable_ocr
                )

            self._converter = DocumentConverter(
                format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=opts)}
            )
        return self._converter

    def extract(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Return (markdown_text, metadata) for a PDF file."""
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        try:
            t0 = time.time()
            result = self._get_converter().convert(str(file_path))
            elapsed = round(time.time() - t0, 2)
            doc = result.document
            content = doc.export_to_markdown()
            metadata = {
                "title": file_path.name,
                "source": str(file_path),
                "mime_type": "application/pdf",
                "pages": len(doc.pages),
                "pictures": len(doc.pictures),
                "tables": len(doc.tables),
                "processing_time": elapsed,
                "extraction_method": "docling",
            }
            logger.info("Extracted PDF", file=file_path.name, pages=metadata["pages"], elapsed_s=elapsed)
            return content, metadata
        except (FileNotFoundError, ExtractionError):
            raise
        except Exception as e:
            logger.error("PDF extraction failed", file=file_path.name, error=str(e))
            raise ExtractionError(f"Failed to extract PDF '{file_path.name}': {e}") from e


class DocxExtractor:
    """Extracts text from DOCX files using python-docx."""

    def extract(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        try:
            import docx
            document = docx.Document(str(file_path))
            paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
            content = "\n\n".join(paragraphs)
            metadata = {
                "title": file_path.name,
                "source": str(file_path),
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "extraction_method": "python-docx",
            }
            logger.info("Extracted DOCX", file=file_path.name, paragraphs=len(paragraphs))
            return content, metadata
        except (FileNotFoundError, ExtractionError):
            raise
        except Exception as e:
            logger.error("DOCX extraction failed", file=file_path.name, error=str(e))
            raise ExtractionError(f"Failed to extract DOCX '{file_path.name}': {e}") from e


class MarkdownExtractor:
    """Reads .md files as plain text (markdown formatting preserved)."""

    def extract(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        try:
            content = file_path.read_text(encoding="utf-8", errors="replace")
            metadata = {
                "title": file_path.name,
                "source": str(file_path),
                "mime_type": "text/markdown",
                "extraction_method": "plain",
            }
            return content, metadata
        except (FileNotFoundError, ExtractionError):
            raise
        except Exception as e:
            raise ExtractionError(f"Failed to read markdown '{file_path.name}': {e}") from e


class PlainTextExtractor:
    """Reads .txt files as plain text."""

    def extract(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        try:
            content = file_path.read_text(encoding="utf-8", errors="replace")
            metadata = {
                "title": file_path.name,
                "source": str(file_path),
                "mime_type": "text/plain",
                "extraction_method": "plain",
            }
            return content, metadata
        except (FileNotFoundError, ExtractionError):
            raise
        except Exception as e:
            raise ExtractionError(f"Failed to read text file '{file_path.name}': {e}") from e


class HtmlExtractor:
    """Extracts readable text from HTML files using trafilatura."""

    def extract(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        try:
            import trafilatura
            raw_html = file_path.read_text(encoding="utf-8", errors="replace")
            # trafilatura.extract returns None if it can't parse; fall back to raw html
            content = trafilatura.extract(raw_html) or raw_html
            metadata = {
                "title": file_path.name,
                "source": str(file_path),
                "mime_type": "text/html",
                "extraction_method": "trafilatura",
            }
            logger.info("Extracted HTML", file=file_path.name, content_len=len(content))
            return content, metadata
        except (FileNotFoundError, ExtractionError):
            raise
        except Exception as e:
            logger.error("HTML extraction failed", file=file_path.name, error=str(e))
            raise ExtractionError(f"Failed to extract HTML '{file_path.name}': {e}") from e


class ImageExtractor:
    """
    Extracts two complementary text representations from image files:

      description  — generated by SmolVLM (via Docling), captures visual semantics.
                     Works well for photos, diagrams, and charts where OCR is sparse.
      ocr_text     — extracted by pytesseract, captures embedded text verbatim.
                     Works well for screenshots, scanned documents, and table images.

    The pipeline stores each as a separate chunk so that both can participate
    in vector retrieval independently.  Non-multimodal LLMs use both texts as
    context; multimodal LLMs also receive the raw image bytes at inference time.
    """

    _MIME: Dict[str, str] = {
        ".jpg":  "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png":  "image/png",
        ".gif":  "image/gif",
        ".webp": "image/webp",
        ".bmp":  "image/bmp",
    }

    def __init__(self, enable_ocr: bool = True, enable_description: bool = True):
        self._enable_ocr = enable_ocr
        self._enable_description = enable_description
        self._converter = None   # Docling converter, lazy-init
        self._converter_failed = False  # skip retry after init failure

    # ── VLM description via Docling ──────────────────────────────────────────

    def _get_converter(self):
        if not self._enable_description or self._converter_failed:
            return None
        if self._converter is not None:
            return self._converter

        try:
            from docling.document_converter import DocumentConverter
            from docling.datamodel.base_models import InputFormat
            from docling.datamodel.pipeline_options import PdfPipelineOptions

            opts = PdfPipelineOptions()
            opts.do_ocr = False               # pytesseract handles OCR
            opts.do_picture_description = True # SmolVLM generates the description
            opts.do_table_structure = False
            opts.images_scale = 2.0

            # docling 2.x exposes ImageFormatOption; fall back to PdfFormatOption
            try:
                from docling.document_converter import ImageFormatOption
                fmt_opt = ImageFormatOption(pipeline_options=opts)
            except ImportError:
                from docling.document_converter import PdfFormatOption
                fmt_opt = PdfFormatOption(pipeline_options=opts)

            self._converter = DocumentConverter(
                format_options={InputFormat.IMAGE: fmt_opt}
            )
            logger.info("Docling image converter ready (SmolVLM)")
        except Exception as e:
            logger.warning(
                "Docling image converter unavailable — descriptions disabled",
                error=str(e),
            )
            self._converter_failed = True

        return self._converter

    def _describe(self, file_path: Path) -> str:
        converter = self._get_converter()
        if converter is None:
            return ""
        try:
            result = converter.convert(str(file_path))
            doc = result.document

            # Primary: look for SmolVLM annotations on picture items
            for pic in doc.pictures:
                for ann in getattr(pic, "annotations", []):
                    text = getattr(ann, "text", "").strip()
                    if text:
                        return text

            # Fallback: use the full markdown export (description may be embedded there)
            md = doc.export_to_markdown().strip()
            return md if md else ""
        except Exception as e:
            logger.warning("Image description failed", file=file_path.name, error=str(e))
            return ""

    # ── OCR via pytesseract ──────────────────────────────────────────────────

    def _ocr(self, file_path: Path) -> str:
        if not self._enable_ocr:
            return ""
        try:
            import pytesseract
            from PIL import Image as PILImage
            from neat_rag.config import settings
            img = PILImage.open(file_path)
            try:
                return pytesseract.image_to_string(img, lang=settings.OCR_LANGUAGES).strip()
            except pytesseract.pytesseract.TesseractError:
                # Some configured language packs may not be installed locally;
                # fall back to English-only so OCR still produces a result.
                logger.warning(
                    "OCR failed with configured languages, falling back to English",
                    file=file_path.name,
                    configured=settings.OCR_LANGUAGES,
                )
                return pytesseract.image_to_string(img, lang="eng").strip()
        except Exception as e:
            logger.warning("OCR failed", file=file_path.name, error=str(e))
            return ""

    # ── Public interface ─────────────────────────────────────────────────────

    def extract(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """
        Returns (primary_content, metadata).

        primary_content  — description if available, else OCR text.
                           Used only for the pipeline's non-empty guard; the
                           pipeline reads description and ocr_text from metadata
                           to build the two separate chunks.
        metadata keys:
          is_image         bool  — signals the pipeline to use image mode
          description      str   — VLM description (may be empty)
          ocr_text         str   — pytesseract output (may be empty)
          mime_type        str
          processing_time  float
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = file_path.suffix.lower()
        mime_type = self._MIME.get(ext, "image/jpeg")

        try:
            t0 = time.time()
            description = self._describe(file_path)
            ocr_text = self._ocr(file_path)
            elapsed = round(time.time() - t0, 2)

            primary_content = description or ocr_text
            if not primary_content:
                raise ExtractionError(
                    f"Image '{file_path.name}' produced neither a VLM description nor OCR "
                    "text. Check that pytesseract is installed and the image is readable."
                )

            metadata: Dict[str, Any] = {
                "title": file_path.name,
                "source": str(file_path),
                "mime_type": mime_type,
                "is_image": True,
                "description": description,
                "ocr_text": ocr_text,
                "processing_time": elapsed,
                "extraction_method": "docling+pytesseract",
            }
            logger.info(
                "Image extracted",
                file=file_path.name,
                has_description=bool(description),
                has_ocr=bool(ocr_text),
                elapsed_s=elapsed,
            )
            return primary_content, metadata

        except (FileNotFoundError, ExtractionError):
            raise
        except Exception as e:
            logger.error("Image extraction failed", file=file_path.name, error=str(e))
            raise ExtractionError(f"Failed to extract image '{file_path.name}': {e}") from e


# ---------------------------------------------------------------------------
# Dispatch table — maps file extension → extractor class
# ---------------------------------------------------------------------------

_EXT_MAP: Dict[str, type] = {
    ".pdf":  PdfExtractor,
    ".docx": DocxExtractor,
    ".md":   MarkdownExtractor,
    ".txt":  PlainTextExtractor,
    ".html": HtmlExtractor,
    ".htm":  HtmlExtractor,
    # Image formats — processed by ImageExtractor (VLM description + OCR)
    ".jpg":  ImageExtractor,
    ".jpeg": ImageExtractor,
    ".png":  ImageExtractor,
    ".gif":  ImageExtractor,
    ".webp": ImageExtractor,
    ".bmp":  ImageExtractor,
}

SUPPORTED_EXTENSIONS = set(_EXT_MAP.keys())

# Extractor singletons — heavy extractors (PdfExtractor/Docling) are expensive
# to initialize; reuse instances across requests to avoid reloading models.
_extractor_cache: Dict[str, Any] = {}


def dispatch_by_ext(file_path: Path) -> "PdfExtractor | DocxExtractor | MarkdownExtractor | PlainTextExtractor | HtmlExtractor":
    """Return a cached extractor instance for the given file extension."""
    ext = file_path.suffix.lower()
    extractor_cls = _EXT_MAP.get(ext)
    if extractor_cls is None:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{ext}'. Supported: {sorted(SUPPORTED_EXTENSIONS)}"
        )
    if ext not in _extractor_cache:
        _extractor_cache[ext] = extractor_cls()
    return _extractor_cache[ext]
